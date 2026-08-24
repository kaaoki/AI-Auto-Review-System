"""
文書自動レビューツール（プロトタイプ：複数フォーマット・フォルダ一括対応）
----------------------------------------
指定フォルダ内の複数ファイル（.txt / .md / .docx / .xlsx / .pdf）を読み込み、
Gemini APIに固定プロンプトでレビューを依頼し、指摘結果をJSON構造化データとして取得。
全ファイルの結果を1つのExcelブックに、ファイルごとのシートとしてまとめて出力する。

事前準備 (Windowsコマンドプロンプトで実行):
    setx GEMINI_API_KEY "APIキー"
    pip install google-genai openpyxl python-docx pdfplumber
"""

import json
import os
import re
import sys
from datetime import datetime

import pdfplumber
from docx import Document as DocxDocument
from google import genai
from google.genai import types
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

MODEL_NAME = "gemini-3.6-flash"

SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".xlsx", ".xlsm", ".pdf"}

#Geminiへのプロンプト
REVIEW_PROMPT_TEMPLATE = """\
あなたは業務文書のレビュー担当者です。以下の文書をレビューし、
次の観点で問題点を洗い出してください。

- 誤字脱字
- 日付・数値など事実関係の誤り（存在しない日付、計算の矛盾など）
- 文章内での矛盾（例：良い結果と悪い結果を並べているのに、片方の評価しか書いていない／前後の文で結論が食い違っている／数値の増減の説明と実際の数値が合っていない）
- フォーマット・表記ゆれ（敬語の統一、表記の一貫性など）

各指摘について、該当箇所の抜粋・指摘内容・修正提案・重要度（高/中/低）を
必ず日本語で、以下のJSON形式のみで返してください。説明文やコードブロックの
```json のような装飾は不要です。JSON以外の文字列を含めないでください。

{{
  "issues": [
    {{
      "excerpt": "該当箇所の抜粋",
      "problem": "指摘内容",
      "suggestion": "修正提案",
      "severity": "高" | "中" | "低"
    }}
  ]
}}

問題が見つからない場合は "issues": [] を返してください。

--- レビュー対象文書 ---
{document_text}
--- 文書ここまで ---
"""


# ---------------------------------------------------------------------------
# ファイル読み込み（フォーマットごとにテキスト化）
# ---------------------------------------------------------------------------

def load_txt_or_md(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def load_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表の中の文字列も拾っておく
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_xlsx(file_path: str) -> str:
    wb = load_workbook(file_path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[シート: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None and str(cell).strip() for cell in row):
                line = "\t".join("" if cell is None else str(cell) for cell in row)
                parts.append(line)
    return "\n".join(parts)


def load_pdf(file_path: str) -> str:
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"[{page_num}ページ]\n{text}")
    return "\n".join(parts)


LOADERS = {
    ".txt": load_txt_or_md,
    ".md": load_txt_or_md,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
    ".xlsm": load_xlsx,
    ".pdf": load_pdf,
}


def load_document(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"未対応のファイル形式です: {ext}")
    return loader(file_path)


# ---------------------------------------------------------------------------
# API呼び出し
# ---------------------------------------------------------------------------

def review_document(document_text: str, client: genai.Client) -> dict:
    prompt = REVIEW_PROMPT_TEMPLATE.format(document_text=document_text)

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",  # JSON出力を強制
        ),
    )
    return json.loads(response.text)


# ---------------------------------------------------------------------------
# コンソール表示
# ---------------------------------------------------------------------------

def print_report(result: dict, file_name: str) -> None:
    issues = result.get("issues", [])
    print(f"\n=== レビュー結果: {file_name} ===")
    if not issues:
        print("指摘事項は見つかりませんでした。")
        return

    for i, issue in enumerate(issues, start=1):
        print(f"\n[{i}] 重要度: {issue.get('severity', '-')}")
        print(f"  該当箇所 : {issue.get('excerpt', '-')}")
        print(f"  指摘内容 : {issue.get('problem', '-')}")
        print(f"  修正提案 : {issue.get('suggestion', '-')}")

    print(f"\n合計 {len(issues)} 件の指摘")


# ---------------------------------------------------------------------------
# Excel出力（全ファイル分をまとめて1ブックに）
# ---------------------------------------------------------------------------

SEVERITY_FILL = {
    "高": PatternFill(start_color="F4CCCC", end_color="F4CCCC", fill_type="solid"),
    "中": PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid"),
    "低": PatternFill(start_color="EFEFEF", end_color="EFEFEF", fill_type="solid"),
}

HEADER_FONT = Font(name="Arial", bold=True, color="FFFFFF", size=11)
HEADER_FILL = PatternFill(start_color="028090", end_color="028090", fill_type="solid")
BODY_FONT = Font(name="Arial", size=10)
THIN_BORDER = Border(
    left=Side(style="thin", color="CCCCCC"),
    right=Side(style="thin", color="CCCCCC"),
    top=Side(style="thin", color="CCCCCC"),
    bottom=Side(style="thin", color="CCCCCC"),
)
WRAP_TOP = Alignment(wrap_text=True, vertical="top")


def _safe_sheet_name(name: str, used_names: set) -> str:
    """Excelのシート名制限（31文字・使用不可文字なし・重複不可）に対応させる"""
    cleaned = re.sub(r"[\[\]:*?/\\]", "_", name)[:31]
    candidate = cleaned or "Sheet"
    suffix = 1
    while candidate in used_names:
        trimmed = cleaned[: 31 - len(f"_{suffix}")]
        candidate = f"{trimmed}_{suffix}"
        suffix += 1
    used_names.add(candidate)
    return candidate


def _write_issue_sheet(ws, file_name: str, result: dict) -> None:
    issues = result.get("issues", []) if "error" not in result else []

    ws["A1"] = f"対象ファイル： {file_name}"
    ws["A2"] = f"レビュー日時： {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ws.merge_cells("A1:E1")
    ws.merge_cells("A2:E2")
    for cell in ("A1", "A2"):
        ws[cell].font = Font(name="Arial", bold=True, size=10)

    if "error" in result:
        ws["A3"] = f"※ このファイルはレビュー処理中にエラーが発生しました: {result['error']}"
        ws["A3"].font = Font(name="Arial", size=10, color="B23A3A")
        ws.merge_cells("A3:E3")

    header_row = 4
    headers = ["No", "重要度", "該当箇所", "指摘内容", "修正提案"]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    for i, issue in enumerate(issues, start=1):
        row = header_row + i
        severity = issue.get("severity", "-")
        values = [
            i,
            severity,
            issue.get("excerpt", "-"),
            issue.get("problem", "-"),
            issue.get("suggestion", "-"),
        ]
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.font = BODY_FONT
            cell.border = THIN_BORDER
            cell.alignment = WRAP_TOP
            if col_idx == 2 and severity in SEVERITY_FILL:
                cell.fill = SEVERITY_FILL[severity]

    if not issues and "error" not in result:
        ws.cell(row=header_row + 1, column=1, value="指摘事項は見つかりませんでした。").font = BODY_FONT

    widths = {"A": 6, "B": 8, "C": 45, "D": 40, "E": 40}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    for row in range(header_row + 1, header_row + 1 + max(len(issues), 1)):
        ws.row_dimensions[row].height = 60

    ws.freeze_panes = f"A{header_row + 1}"


def _write_summary_sheet(ws, results: dict) -> None:
    ws["A1"] = "レビュー対象ファイル一覧（サマリー）"
    ws["A1"].font = Font(name="Arial", bold=True, size=13)
    ws.merge_cells("A1:D1")

    header_row = 3
    headers = ["ファイル名", "指摘件数", "うち重要度:高", "ステータス"]
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = THIN_BORDER

    for i, (file_name, result) in enumerate(results.items(), start=1):
        row = header_row + i
        if "error" in result:
            values = [file_name, "-", "-", f"エラー: {result['error']}"]
        else:
            issues = result.get("issues", [])
            high_count = sum(1 for iss in issues if iss.get("severity") == "高")
            values = [file_name, len(issues), high_count, "完了"]
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.font = BODY_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = {"A": 35, "B": 12, "C": 14, "D": 30}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width


def write_excel_report(results: dict, output_path: str) -> None:
    """
    results: { "ファイル名": { "issues": [...] } または { "error": "..." } } の辞書
    """
    wb = Workbook()

    summary_ws = wb.active
    summary_ws.title = "サマリー"
    _write_summary_sheet(summary_ws, results)

    used_names = {"サマリー"}
    for file_name, result in results.items():
        sheet_name = _safe_sheet_name(os.path.splitext(file_name)[0], used_names)
        ws = wb.create_sheet(title=sheet_name)
        _write_issue_sheet(ws, file_name, result)

    wb.save(output_path)
    print(f"\nExcelレポートを出力しました: {output_path}")


# ---------------------------------------------------------------------------
# メイン処理
# ---------------------------------------------------------------------------

def collect_target_files(input_path: str) -> list:
    if os.path.isfile(input_path):
        return [input_path]

    files = []
    for name in sorted(os.listdir(input_path)):
        full_path = os.path.join(input_path, name)
        if os.path.isfile(full_path) and os.path.splitext(name)[1].lower() in SUPPORTED_EXTENSIONS:
            files.append(full_path)
    return files


def main():
    if len(sys.argv) < 2:
        print("使い方: python review_tool.py <フォルダ or ファイルパス> [出力Excelパス]")
        print(f"対応形式: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        sys.exit(1)

    input_path = sys.argv[1]
    if not os.environ.get("GEMINI_API_KEY"):
        print("エラー: 環境変数 GEMINI_API_KEY が設定されていません。")
        sys.exit(1)

    target_files = collect_target_files(input_path)
    if not target_files:
        print(f"対象ファイルが見つかりませんでした: {input_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(input_path)) or ".", "output")
        os.makedirs(output_dir, exist_ok=True)
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = os.path.join(output_dir, f"review_report_{stamp}.xlsx")

    client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

    results = {}
    print(f"{len(target_files)} 件のファイルをレビューします...")
    for file_path in target_files:
        file_name = os.path.basename(file_path)
        print(f"\n--- 処理中: {file_name} ---")
        try:
            document_text = load_document(file_path)
            if not document_text.strip():
                results[file_name] = {"error": "テキストを抽出できませんでした（空ファイルの可能性）"}
                continue
            result = review_document(document_text, client)
            results[file_name] = result
            print_report(result, file_name)
        except Exception as e:  # ファイル単位のエラーで全体を止めない
            print(f"エラーが発生しました: {e}")
            results[file_name] = {"error": str(e)}

    write_excel_report(results, output_path)


if __name__ == "__main__":
    main()
